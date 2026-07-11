from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path

OUT = Path('output')
OUT.mkdir(exist_ok=True)
DOCX = OUT / 'L15_L16_Neural_Networks_Lecture_Note_CN.docx'

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 100, 112)
LIGHT_BLUE = 'E8EEF5'
PALE = 'F4F7FA'
GOLD = RGBColor(176, 124, 32)
BLACK = RGBColor(0, 0, 0)

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.82)
sec.bottom_margin = Inches(0.78)
sec.left_margin = Inches(0.9)
sec.right_margin = Inches(0.9)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn('w:' + m))
        if node is None:
            node = OxmlElement('w:' + m)
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)


def set_run_font(run, name='Microsoft YaHei', size=None, bold=None, italic=None, color=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Microsoft YaHei'
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.22

for name, size, color, before, after in [
    ('Title', 28, DARK_BLUE, 0, 8),
    ('Subtitle', 13, MUTED, 0, 12),
    ('Heading 1', 16, BLUE, 17, 8),
    ('Heading 2', 13, BLUE, 13, 6),
    ('Heading 3', 11.5, DARK_BLUE, 9, 4),
]:
    s = styles[name]
    s.font.name = 'Microsoft YaHei'
    s._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    s.font.size = Pt(size)
    s.font.color.rgb = color
    s.font.bold = name not in ('Subtitle',)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

if 'Equation' not in styles:
    eq_style = styles.add_style('Equation', WD_STYLE_TYPE.PARAGRAPH)
else:
    eq_style = styles['Equation']
eq_style.font.name = 'Cambria Math'
eq_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Cambria Math')
eq_style.font.size = Pt(10.5)
eq_style.paragraph_format.space_before = Pt(4)
eq_style.paragraph_format.space_after = Pt(5)
eq_style.paragraph_format.keep_together = True

for list_style in ['List Bullet', 'List Number']:
    s = styles[list_style]
    s.font.name = 'Microsoft YaHei'
    s._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    s.font.size = Pt(10.5)
    s.paragraph_format.space_after = Pt(3)
    s.paragraph_format.line_spacing = 1.18


def add_field(paragraph, field):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.extend([fldChar1, instrText, fldChar2])


header = sec.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hr = hp.add_run('MATHEMATICAL FOUNDATIONS OF ML  |  L15–L16 INTEGRATED NOTES')
set_run_font(hr, size=8.3, bold=True, color=MUTED)

footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run('Neural Networks: Representation, Memorization & Optimization   •   ')
set_run_font(fr, size=8, color=MUTED)
add_field(fp, 'PAGE')


def para(text='', bold_prefix=None, italic=False, color=None, align=None, after=None, keep=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_together = keep
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, italic=italic, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic, color=color)
    return p


def equation(text, label=None):
    p = doc.add_paragraph(style='Equation')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    r = p.add_run(text)
    set_run_font(r, name='Cambria Math', size=10.5)
    if label:
        r2 = p.add_run('    ' + label)
        set_run_font(r2, name='Cambria Math', size=9.5, color=MUTED)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    r = p.add_run(text)
    set_run_font(r)
    return p


def restart_list_numbering(paragraph, start=1):
    style = paragraph.style
    num_id = style.element.pPr.numPr.numId.val
    numbering = paragraph.part.numbering_part.element
    old_num = numbering.xpath(f'./w:num[@w:numId="{num_id}"]')[0]
    abstract_id = old_num.abstractNumId.val
    existing = [int(n.get(qn('w:numId'))) for n in numbering.findall(qn('w:num'))]
    new_id = max(existing) + 1
    new_num = OxmlElement('w:num')
    new_num.set(qn('w:numId'), str(new_id))
    abs_el = OxmlElement('w:abstractNumId')
    abs_el.set(qn('w:val'), str(abstract_id))
    new_num.append(abs_el)
    lvl_override = OxmlElement('w:lvlOverride')
    lvl_override.set(qn('w:ilvl'), '0')
    start_override = OxmlElement('w:startOverride')
    start_override.set(qn('w:val'), str(start))
    lvl_override.append(start_override)
    new_num.append(lvl_override)
    numbering.append(new_num)
    pPr = paragraph._p.get_or_add_pPr()
    numPr = pPr.get_or_add_numPr()
    numPr.get_or_add_ilvl().val = 0
    numPr.get_or_add_numId().val = new_id


def number(text, restart=False):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(text)
    set_run_font(r)
    if restart:
        restart_list_numbering(p, 1)
    return p


def callout(title, body, fill=PALE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 110, 160, 110, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + '  ')
    set_run_font(r, bold=True, color=DARK_BLUE)
    r = p.add_run(body)
    set_run_font(r)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_table(headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    for i, (h, w) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[i]
        cell.width = Inches(w)
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=9.3, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, (val, w) in enumerate(zip(row, widths)):
            cells[i].width = Inches(w)
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            set_run_font(r, size=9.1)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def source(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run('来源：' + text)
    set_run_font(r, size=8.2, italic=True, color=MUTED)


# Cover
doc.add_paragraph().paragraph_format.space_after = Pt(54)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('LECTURE NOTE')
set_run_font(r, size=11, bold=True, color=GOLD)
p.paragraph_format.space_after = Pt(14)
p = doc.add_paragraph(style='Title')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('神经网络：表达、记忆与优化')
set_run_font(r, size=28, bold=True, color=DARK_BLUE)
p = doc.add_paragraph(style='Subtitle')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('L15–L16 幻灯片与两本机器学习理论教材的整合讲义')
set_run_font(r, size=13, color=MUTED)

doc.add_paragraph().paragraph_format.space_after = Pt(30)
callout('核心主线', '神经网络同时是一种可学习特征映射、函数逼近器和可优化的复合模型。理解它需要把表达能力、有限样本记忆、优化动力学与泛化控制放到同一张图里。', 'EDF3F8')

doc.add_paragraph().paragraph_format.space_after = Pt(30)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('课程：CIT413048 Mathematical Foundations of ML\n')
set_run_font(r, size=10.2, color=MUTED)
r = p.add_run('基于 Suvrit Sra 的 L15 / L16（2026-06-09）\n')
set_run_font(r, size=10.2, color=MUTED)
r = p.add_run('补充：Francis Bach, Learning Theory from First Principles；\nShalev-Shwartz & Ben-David, Understanding Machine Learning')
set_run_font(r, size=9.6, color=MUTED)

doc.add_page_break()

doc.add_heading('阅读导航', level=1)
para('这份讲义不是逐页复述，而是围绕“模型是什么—能表示什么—如何训练—为何可能泛化”重排材料。公式采用统一的列向量约定，避免原幻灯片与教材之间转置方向不一致。')
add_table(
    ['模块', '要回答的问题', '主要材料'],
    [
        ('A. 表示', '网络如何从固定特征走向可学习特征？', 'L15 pp. 2–25；Bach §9.1–9.2'),
        ('B. 表达能力', '万能逼近说明了什么，又没有说明什么？', 'L15 pp. 26–40；Bach §9.3–9.5；UML §20.3–20.4'),
        ('C. 记忆', '有限样本插值与 VC 维有什么区别？', 'L16 pp. 3–14'),
        ('D. 优化', 'SGD 与反向传播如何得到梯度？', 'L16 pp. 15–36；Bach §5.4；UML §20.6'),
        ('E. 稳定与正则化', '初始化、Dropout、BN 如何影响训练？', 'L16 pp. 20–47'),
        ('F. 泛化', '参数多为何不必然意味着泛化差？', 'Bach §9.2.3–9.4；UML §20.4'),
    ],
    [1.15, 3.15, 2.2],
)

doc.add_heading('符号约定', level=2)
add_table(
    ['符号', '含义'],
    [
        ('x ∈ ℝᵈ, y', '输入与标签'),
        ('L', '不计输入层的网络深度'),
        ('Wˡ ∈ ℝ^{mₗ×mₗ₋₁}, bˡ ∈ ℝ^{mₗ}', '第 l 层权重与偏置'),
        ('zˡ, aˡ', '预激活与激活，a⁰=x'),
        ('σ', '逐坐标激活函数；ReLU(u)=max{u,0}'),
        ('θ', '所有权重和偏置的集合'),
        ('R(θ), R̂ₙ(θ)', '总体风险与经验风险'),
    ],
    [1.7, 4.8],
)

doc.add_heading('1. 从固定特征到可学习特征', level=1)
doc.add_heading('1.1 核方法视角', level=2)
para('核方法先选定核 k，再由训练数据产生隐式特征。以核 SVM 为例，表示定理给出 w=∑ᵢ αᵢyᵢφ(xᵢ)，预测为：')
equation('⟨w, φ(x)⟩ = ∑ᵢ αᵢ yᵢ k(xᵢ,x) = ⟨α, Φ(x)⟩,    Φ(x)=[y₁k(x₁,x),…,yₙk(xₙ,x)]ᵀ.')
para('因此分类器对 Φ(x) 是线性的，但 Φ 的形状由核预先固定。神经网络的关键变化是：把特征本身参数化，并与最后的线性分类器联合学习。')
equation('x  ↦  Φ_{θ₁}(x)  ↦  wᵀΦ_{θ₁}(x)+b,    θ=(θ₁,w,b).')
callout('四种等价直觉', '神经网络可被看成：①可学习特征；②嵌套分类器；③由简单决策逐层组成复杂决策；④可组合的函数逼近器。', 'EEF4F8')
source('L15 pp. 5–13；Bach, Ch. 9, pp. 247–250。')

doc.add_heading('1.2 前馈网络的统一写法', level=2)
para('令 a⁰=x。对 l=1,…,L，定义')
equation('zˡ = Wˡaˡ⁻¹+bˡ,      aˡ = σˡ(zˡ).', '(1.1)')
para('若最后一层为线性输出，则 F(x;θ)=zᴸ；若是分类问题，常把 zᴸ 解释为 logits，再配合 sigmoid 或 softmax 与交叉熵。输入层只存储数据，不计入深度。')
equation('F(x;θ)=Wᴸ σ^{L−1}(W^{L−1}⋯σ¹(W¹x+b¹)⋯+b^{L−1})+bᴸ.')

doc.add_heading('1.3 一层与单隐层网络', level=2)
para('没有隐藏层时，网络退化为线性模型 z=wᵀx+b；加上 sign 得到线性分类器，加上 sigmoid 与对数损失得到逻辑回归。单隐层标量输出网络为')
equation('f(x)=∑_{j=1}^m ηⱼ σ(wⱼᵀx+bⱼ).', '(1.2)')
para('固定 (wⱼ,bⱼ) 时，这是在 m 个固定特征 σ(wⱼᵀx+bⱼ) 上的线性模型；同时优化输入权重时，模型开始“学习特征”。')
source('L15 pp. 14–18；Bach §9.2；UML §20.1。')

doc.add_heading('2. 激活函数、输出层与损失', level=1)
add_table(
    ['函数', '定义', '导数/性质', '典型作用'],
    [
        ('Sigmoid', 'σ(u)=1/(1+e⁻ᵘ)', 'σ′=σ(1−σ)，大 |u| 饱和', '二分类概率输出'),
        ('tanh', 'tanh(u)', '1−tanh²(u)，零中心但会饱和', '传统隐藏层'),
        ('ReLU', '[u]₊=max(0,u)', '1{u>0}（u=0 取次梯度）', '现代隐藏层基础'),
        ('Softmax', 'pₖ=e^{zₖ}/∑ⱼe^{zⱼ}', 'J=diag(p)−ppᵀ', '多分类概率输出'),
    ],
    [1.0, 1.8, 2.0, 1.7],
)

doc.add_heading('2.1 交叉熵与逻辑损失的等价性', level=2)
para('对二分类 y∈{−1,+1}，令 g(x)=sigmoid(f(x))。负对数似然为')
equation('−1{y=1}log g(x) − 1{y=−1}log(1−g(x)) = log(1+e^{−yf(x)}).', '(2.1)')
para('所以“sigmoid + binary cross-entropy”与直接对 logit 使用 logistic loss 完全等价。数值实现通常直接从 logits 计算，以避免先取概率产生上溢或下溢。')

doc.add_heading('2.2 Softmax 交叉熵梯度', level=2)
para('若真实类别的 one-hot 向量为 y，p=softmax(z)，损失 ℓ=−∑ₖyₖlog pₖ，则')
equation('∂ℓ/∂z = p−y.', '(2.2)')
para('推导：∂log pᵣ/∂zₖ=δᵣₖ−pₖ，因此 ∂ℓ/∂zₖ=−∑ᵣyᵣ(δᵣₖ−pₖ)=pₖ−yₖ。这个简洁结果是多分类反向传播的起点。')

doc.add_heading('3. 非线性特征如何产生线性可分性', level=1)
para('两个 ReLU 单元定义两个半空间响应：')
equation('z₁=w₁ᵀx+b₁,  z₂=w₂ᵀx+b₂,      Φ(x)=([z₁]₊,[z₂]₊).')
para('原空间中非线性可分的数据，在 Φ(x) 空间中可能被一条直线分开。每个 ReLU 把一个仿射超平面一侧截为零，另一侧保留线性距离信息；多个单元共同把输入空间切成多面体区域，并在每个区域上实现不同的仿射映射。')
callout('重要限制', '表示存在性不等于训练可达性。即使某组超平面能把样本变得线性可分，梯度法仍需从随机初始化中找到合适方向；法向量符号翻转就可能产生完全不同的特征。')
source('L15 pp. 19–25。')

doc.add_heading('4. 表达能力与万能逼近', level=1)
doc.add_heading('4.1 万能逼近定理的准确含义', level=2)
para('经典结果表明：在紧集上，只要激活函数满足适当条件，单隐层网络可以把任意连续函数逼近到任意精度。对 ReLU，非多项式性足以推出稠密性。形式化地，对紧集 K⊂ℝᵈ、连续 g 与任意 ε>0，存在 m 与参数使')
equation('sup_{x∈K} |∑_{j=1}^m ηⱼσ(wⱼᵀx+bⱼ) − g(x)| < ε.', '(4.1)')
callout('万能逼近没有保证的三件事', '它不保证所需宽度小、不保证训练算法能找到参数，也不保证从有限数据学到的函数具有小测试误差。表达、优化与泛化必须分开讨论。', 'FFF7E8')

doc.add_heading('4.2 一维 ReLU 表示的直接推导', level=2)
para('设 f 在 [−R,R] 上连续分段仿射，折点 a₁<⋯<aₖ，各区间斜率为 v₀,…,vₖ。ReLU 的导数是阶跃函数，所以每增加一个 ReLU 就能在一个折点改变斜率。由此得到精确表示')
equation('f(x)=f(−R)+v₀(x+R)+∑_{j=1}^k (vⱼ−vⱼ₋₁)(x−aⱼ)₊.', '(4.2)')
para('验证很直接：在 x<a₁ 时导数为 v₀；越过 aⱼ 后，导数增加 vⱼ−vⱼ₋₁，于是新斜率变为 vⱼ。连续函数可被细网格上的分段线性插值一致逼近，因此得到一维万能逼近。')

doc.add_heading('4.3 光滑函数的积分表示', level=2)
para('若 f 二次连续可微，Taylor 公式的积分余项给出')
equation('f(x)=f(−R)+f′(−R)(x+R)+∫_{−R}^{R} f″(b)(x−b)₊ db.', '(4.3)')
para('这说明 ReLU 不只是离散折线基：连续的 ReLU 混合可以精确重建光滑函数。输出权重总变差由 ∫|f″(b)|db 控制，为后面的范数型泛化界提供桥梁。')
source('Bach §9.3.1–9.3.3, pp. 256–263。')

doc.add_heading('4.4 多维证明框架与三层构造', level=2)
number('利用连续性：把 [0,1]ᵈ 划分为足够小的超矩形，在每个小块内用常数近似 g。')
number('把分段常数写成超矩形指示函数的线性组合。')
number('用两层 ReLU 构造连续“软指示器”，在边界薄层外逼近每个超矩形指示函数。')
number('最后用线性输出层叠加这些软指示器，从而逼近 g。')
para('该思路解释了 L15 中“三层 ReLU 网络万能逼近”的直接证明。注意幻灯片把积分 L¹ 误差写作 ‖·‖∞ 的位置应区分：一致误差是 sup 范数，积分绝对误差是 L¹ 范数。')
source('L15 pp. 28–30；UML Ch. 20, Exercise 1。')

doc.add_heading('4.5 深度为何可能节省宽度', level=2)
para('万能逼近只说“能表示”，而深度分离关注“表示得是否高效”。L15 给出的典型例子是单位球指示函数：三层网络先近似每个坐标平方，再求和得到 ‖x‖²，最后近似阈值；宽度对 d 与 1/ε 为多项式。某些两层网络若要达到固定高精度，则需要指数级宽度。')
equation('x  ↦  (x₁²,…,x_d²)  ↦  ∑ᵢxᵢ²=‖x‖₂²  ↦  1{‖x‖₂≤1}.')
para('本质是复用中间计算。深层网络把可组合结构 f∘g∘h 显式编码；浅层网络可能只能用大量基函数“平铺”同一结构。深度并非总有优势，优势依赖目标函数与架构之间是否存在匹配的组合结构。')
source('L15 pp. 31–39；Safran & Shamir (2017) 摘要性结果。')

doc.add_heading('5. 近似、估计与优化：三种误差', level=1)
para('经验风险最小化的分析应把三种问题分开：函数类是否足够丰富（近似误差）、样本是否足够（估计误差）、算法是否找到足够好的经验解（优化误差）。')
equation('R(f̂)−inf_f R(f) = approximation + estimation + optimization  （概念分解）.')
para('更具体地，令 𝓕 为模型类、f* 为总体最优函数、f*_{𝓕}=argmin_{f∈𝓕}R(f)，f̂ 为算法输出，则可写成')
equation('R(f̂)−R(f*) = [R(f*_{𝓕})−R(f*)] + [R(f̂_ERM)−R(f*_{𝓕})] + [R(f̂)−R(f̂_ERM)].')
para('实际界会通过经验风险与总体风险的偏差来控制第二项。这个框架能澄清一个常见误区：训练误差为零只意味着优化/插值成功，不能单独推出总体风险小。')

doc.add_heading('5.1 ReLU 网络的范数与 Rademacher 界', level=2)
para('对单隐层 ReLU 网络 f(x)=∑ⱼηⱼ(wⱼᵀx+bⱼ)₊，假设 ‖x‖₂≤R，归一化 ‖wⱼ‖₂²+bⱼ²/R²=1，并约束 ‖η‖₁≤D。若损失对预测是 G-Lipschitz，则经验风险最小化的期望估计误差满足')
equation('E[R(f̂)−inf_{f∈𝓕}R(f)] ≤ 16GDR/√n.', '(5.1)')
para('推导主链如下。先由收缩不等式去掉损失，再用 ℓ₁/ℓ∞ 对偶把对所有神经元的优化化为对单个神经元的上确界，再用 ReLU 的 1-Lipschitz 性与 Cauchy–Schwarz：')
equation('ℛₙ(ℓ∘𝓕) ≤ 2GD·E sup_{‖w‖²+b²/R²=1}|(1/n)∑ᵢεᵢ(wᵀxᵢ+b)|')
equation('≤ 2GD·E[(‖(1/n)∑ᵢεᵢxᵢ‖²+R²((1/n)∑ᵢεᵢ)²)^{1/2}] ≤ 4GDR/√n.')
para('关键结论是：在这个界里，隐藏单元数 m 没有显式出现；真正控制估计误差的是权重范数。参数数量并非完全无关，但“参数多 ⇒ 必然泛化差”不是正确推理。')
source('Bach §9.2.3, Proposition 9.1。')

doc.add_heading('5.2 ReLU 正齐次性与隐含 ℓ₁ 正则', level=2)
para('ReLU 满足 (αu)₊=α(u)₊（α>0）。因此单个神经元 η(wᵀx+b)₊ 在缩放 η→αη、(w,b)→(w,b)/α 后保持不变。若使用平方 ℓ₂ 权重衰减，单元的惩罚可写为')
equation('α²η² + (‖w‖²+b²/R²)/α².')
para('对 α 优化。令 c²=‖w‖²+b²/R²，由 AM–GM 或求导可得 α²=c/|η|，最小值为')
equation('min_{α>0}[α²η²+c²/α²] = 2|η|c.', '(5.2)')
para('在 c=1 的归一化下，这就是输出权重的 ℓ₁ 惩罚，解释了正齐次网络中平方权重衰减与稀疏/变差范数之间的联系。')
source('Bach §9.2.2。')

doc.add_heading('6. 有限样本表达能力与记忆', level=1)
doc.add_heading('6.1 定义与 VC 维的区别', level=2)
para('给定 N 个互异输入 xᵢ∈ℝᵈ 和任意目标 yᵢ∈ℝᵖ，若存在 θ 使 fθ(xᵢ)=yᵢ 对所有 i 成立，则称网络对该数据集可插值。若对任意这样的数据集都成立，则称具有 N 点的普适有限样本表达能力。')
equation('MemCap(𝓕)=max{N: ∀{(xᵢ,yᵢ)}ᵢ₌₁ᴺ, ∃θ, fθ(xᵢ)=yᵢ}.')
add_table(
    ['概念', '量词结构', '标签范围', '含义'],
    [
        ('记忆容量', '对所有输入点集、所有标签', '通常实数或向量', '最坏数据也能精确插值'),
        ('VC 维', '存在一个输入点集，可实现所有二元标记', '{−1,+1}', '至少有一个点集可被打散'),
    ],
    [1.1, 2.2, 1.2, 2.0],
)
equation('Memorization capacity ≤ VC dimension  （标量二分类情形）.')

doc.add_heading('6.2 深度—宽度权衡', level=2)
para('L16 汇总的结果（在互异输入、分段线性激活等条件下）包括：')
bullet('单隐层 ReLU 网络用 Θ(N) 个隐藏单元可以记忆 N 个任意样本。')
bullet('三隐层充分条件：d₁d₂≥4N 且 d₃≥4p；两隐层可在 d₁d₂≥4Np 时记忆 N 点 p 维输出。')
bullet('必要性示例：单隐层若 d₁+2<N，或两隐层标量输出若 2d₁d₂+d₂+2<N，则存在无法记忆的 N 点数据集。')
bullet('再增加一层时，多类问题的节点需求可从 Θ(Np) 改善到 Θ(N+p)，体现深度对有限样本表达的节省。')
callout('记忆不等于学习失败', '记忆随机标签说明模型容量很大，也说明零训练误差不足以解释泛化；但真实任务中，优化算法、数据结构、显式正则化与隐式偏置可能在众多插值解中选择更简单的解。')
source('L16 pp. 3–14；Yun, Sra & Jadbabaie (2019)。')

doc.add_heading('6.3 VC 维与样本复杂度', level=2)
para('对 sign 激活的固定有向图，UML 给出')
equation('VCdim(H_{V,E,sign}) = O(|E| log|E|).', '(6.1)')
para('证明轮廓：每个神经元是半空间，其 m 点增长函数至多 (em/d)^d；网络是各层与各单元函数类的复合/乘积，因此增长函数相乘，得到 τ_H(m)≤(em)^{|E|}。若 m 点被打散，则 2^m≤(em)^{|E|}，从而 m=O(|E|log|E|)。现代分段线性网络更精细的界还显式依赖深度 L。')
source('UML §20.4；L15 p. 40 引用 Bartlett et al. (2019)。')

doc.add_heading('7. 经验风险与随机梯度下降', level=1)
para('训练目标通常是正则化经验风险')
equation('R̂ₙ(θ)= (1/N)∑_{i=1}^N ℓ(yᵢ,F(xᵢ;θ)) + (λ/2)‖θ‖².', '(7.1)')
para('随机选取 iₜ 或 mini-batch Bₜ，构造梯度估计 gₜ。SGD 更新为')
equation('θₜ₊₁ = θₜ − ηₜ gₜ,      E[gₜ | θₜ]=∇R̂ₙ(θₜ).', '(7.2)')
para('无偏性使噪声项在条件期望下消失。对凸、B-Lipschitz 目标，若 ‖θ*−θ₀‖≤D，取 ηₜ≈D/(B√t)，加权平均迭代点满足 O(DB log t/√t) 的界；若预先知道总步数 T 并用常数步长 η=D/(B√T)，则')
equation('E[F(θ̄_T)−F(θ*)] ≤ D²/(2ηT)+ηB²/2 = DB/√T.', '(7.3)')
para('推导核心来自平方距离递推：')
equation('E‖θₜ₊₁−θ*‖² = E‖θₜ−θ*‖² −2ηₜE⟨∇F(θₜ),θₜ−θ*⟩ +ηₜ²E‖gₜ‖²,')
para('再用凸性 F(θₜ)−F(θ*)≤⟨∇F(θₜ),θₜ−θ*⟩ 并求和望远镜消去距离项。')
callout('非凸警告', '神经网络目标高度非凸，式 (7.3) 的凸收敛保证不能直接套用。一般光滑非凸分析通常只能保证某个迭代点的梯度范数较小，而不是到达全局最优；过参数化等额外结构可带来更强结论。', 'FFF7E8')
source('L16 pp. 15–22；Bach §5.4；UML §20.6。')

doc.add_heading('8. 初始化：为什么 He 方差是 2/n', level=1)
para('零初始化会使同层单元保持完全对称；对 ReLU，某些结构下还会使激活与梯度为零。因此用独立、零均值随机权重破坏对称性。He 初始化的方差可由前向方差守恒近似推出。')
para('设某层有 n 个独立输入 aⱼ，权重 wⱼ 独立、E[wⱼ]=0，忽略偏置：')
equation('z=∑_{j=1}^n wⱼaⱼ  ⇒  Var(z)=n Var(w) E[a²].')
para('若 z 近似关于 0 对称，ReLU 保留约一半二阶矩：E[ReLU(z)²]≈(1/2)E[z²]。希望下一层的二阶矩与当前层相当，则')
equation('E[a_{next}²] ≈ (1/2)n Var(w)E[a²] ≈ E[a²]  ⇒  Var(w)=2/n.', '(8.1)')
para('因此 wⱼ∼N(0,2/n) 或同方差的均匀分布。对线性/tanh 型激活常得到 Xavier 尺度约 1/n；更细致的版本同时考虑 fan-in 与 fan-out。')
source('L16 pp. 20–21；He et al. (2015)。')

doc.add_heading('9. 反向传播：矩阵推导', level=1)
doc.add_heading('9.1 前向与误差信号', level=2)
para('采用式 (1.1) 的列向量约定。定义误差信号 δˡ=∂ℓ/∂zˡ。输出层先由损失确定 δᴸ；例如线性输出配平方损失时 δᴸ=aᴸ−y，softmax 交叉熵时 δᴸ=p−y。')
equation('δᴸ = ∇_{zᴸ}ℓ.', '(9.1)')

doc.add_heading('9.2 递推公式', level=2)
para('因为 zˡ⁺¹=Wˡ⁺¹aˡ+bˡ⁺¹，aˡ=σ(zˡ)，链式法则给出')
equation('∂ℓ/∂aˡ = (Wˡ⁺¹)ᵀδˡ⁺¹,')
equation('δˡ = [(Wˡ⁺¹)ᵀδˡ⁺¹] ⊙ σ′(zˡ),      l=L−1,…,1.', '(9.2)')
para('等价矩阵形式是 δˡ=Diag(σ′(zˡ))(Wˡ⁺¹)ᵀδˡ⁺¹。原幻灯片使用 W 的转置定义，因而公式外观不同；二者只要维度一致就是同一递推。')

doc.add_heading('9.3 参数梯度', level=2)
para('对单个元素 zᶫᵢ=∑ⱼWᶫᵢⱼaᶫ⁻¹ⱼ+bᶫᵢ，有 ∂zᶫᵢ/∂Wᶫᵢⱼ=aᶫ⁻¹ⱼ，因此')
equation('∂ℓ/∂Wˡ = δˡ(aˡ⁻¹)ᵀ,      ∂ℓ/∂bˡ=δˡ.', '(9.3)')
para('反向传播的效率来自动态规划：前向保存每层 zˡ,aˡ；反向每条边只参与常数次乘加，计算全部参数梯度的时间与一次前向计算同阶，而不是对每个参数重新跑一次链式法则。')

doc.add_heading('9.4 单隐层 hinge-loss 示例', level=2)
para('设 zᵢ=wᵢᵀx+bᵢ，hᵢ=[zᵢ]₊，s=∑ᵢvᵢhᵢ+c，ℓ(y,s)=max(0,1−ys)。则')
equation('∂ℓ/∂s = −y·1{1−ys>0},')
equation('∂ℓ/∂wᵢ = (−y·1{1−ys>0}) vᵢ 1{zᵢ>0} x.', '(9.4)')
para('这正是 L16 p. 24 的四段链式乘积：输入 x、ReLU 门、输出权重 vᵢ、损失边际门。')
source('L16 pp. 23–33；UML §20.6, pp. 277–281。')

doc.add_heading('10. 梯度消失与爆炸', level=1)
para('把递推展开，可见早期层梯度包含长矩阵乘积：')
equation('δˡ = Dˡ(Wˡ⁺¹)ᵀDˡ⁺¹(Wˡ⁺²)ᵀ⋯Dᴸ δ_out,    Dʳ=Diag(σ′(zʳ)).')
para('取算子范数得到')
equation('‖δˡ‖ ≤ [∏_{r=l}^{L} ‖Dʳ‖] [∏_{r=l+1}^{L} ‖Wʳ‖] ‖δ_out‖.', '(10.1)')
para('若多数因子范数小于 1，乘积指数衰减；若多数大于 1，则指数增长。Sigmoid 饱和区 σ′≈0 会加剧消失；ReLU 的激活侧导数为 1，但仍不能单独解决权重矩阵连乘问题。')
add_table(
    ['方法', '主要作用机制'],
    [
        ('He/Xavier 初始化', '让前向激活与反向梯度的初始尺度近似守恒'),
        ('梯度裁剪', '直接限制爆炸梯度的范数'),
        ('正交/规范化参数化', '控制矩阵的奇异值尺度'),
        ('残差连接', '引入恒等路径，使 Jacobian 含 I 项'),
        ('Batch/Layer Normalization', '稳定中间激活尺度并改变优化几何'),
    ],
    [1.55, 4.95],
)
source('L16 pp. 36–37。')

doc.add_heading('11. 正则化与 Dropout', level=1)
doc.add_heading('11.1 权重衰减', level=2)
para('显式 ℓ₂ 正则化把目标改为 R̂ₙ(θ)+(λ/2)‖θ‖²，梯度更新为')
equation('θ ← θ − η(∇R̂ₙ(θ)+λθ).')
para('在普通 SGD 中这等价于每步先把权重乘以 1−ηλ，再减去数据梯度；在自适应优化器中，“L₂ 正则”与解耦 weight decay 不一定等价。')

doc.add_heading('11.2 Dropout 的随机计算图', level=2)
para('训练时对激活 a 采样独立掩码 rⱼ∼Bernoulli(q)。现代“inverted dropout”写作')
equation('ã = (r⊙a)/q,      E[ã|a]=a.', '(11.1)')
para('因此测试时直接使用 a，无需再缩放。原始表述也可在训练时不除以 q、测试时把输出权重乘 q；两种约定等价。被丢弃单元的局部梯度为零：')
equation('∂ℓ/∂a = (r/q)⊙∂ℓ/∂ã.')
para('Dropout 可理解为训练大量共享参数的“瘦网络”，抑制单元间脆弱共适应，并近似进行模型平均。它增加梯度噪声，保留率 q 需要与网络规模、归一化层和数据量共同调节。')
source('L16 pp. 41–42。')

doc.add_heading('12. Batch Normalization 与完整梯度', level=1)
doc.add_heading('12.1 前向变换', level=2)
para('对一个 mini-batch 中某个标量特征 x₁,…,x_m，BN 定义')
equation('μ_B=(1/m)∑ᵢxᵢ,      σ_B²=(1/m)∑ᵢ(xᵢ−μ_B)²,')
equation('x̂ᵢ=(xᵢ−μ_B)/√(σ_B²+ε),      yᵢ=γx̂ᵢ+β.', '(12.1)')
para('γ 与 β 允许网络恢复或改变尺度与平移。若忽略 ε，取 γ=√Var[x]、β=E[x] 可恢复恒等映射。训练时用批统计量；推理时通常用移动平均的总体统计量。')

doc.add_heading('12.2 反向传播推导结果', level=2)
para('令 dᵢ=∂ℓ/∂yᵢ，s=√(σ_B²+ε)。显然')
equation('∂ℓ/∂β=∑ᵢdᵢ,      ∂ℓ/∂γ=∑ᵢdᵢx̂ᵢ.', '(12.2)')
para('对输入梯度，需要同时考虑 xᵢ 对自身、均值与方差的影响。把链式法则展开并合并，可得到常用紧凑形式')
equation('∂ℓ/∂xᵢ = γ/(m s) · [m dᵢ − ∑ⱼdⱼ − x̂ᵢ∑ⱼ(dⱼx̂ⱼ)].', '(12.3)')
para('式 (12.3) 清楚显示 BN 的梯度在 batch 内耦合：一个样本的梯度取决于整批的 dⱼ 与 x̂ⱼ。也因此不能把均值和方差当作与输入无关的常数。')

doc.add_heading('12.3 尺度不变性与解释边界', level=2)
para('忽略 ε 且 α>0 时，BN(αx)=BN(x)，因此 BN(Wx)=BN(αWx)。这减弱了参数尺度对前向输出的影响，通常允许更大学习率，并对优化产生隐式调节。')
callout('历史解释需谨慎', '“减少 internal covariate shift”是 BN 原论文的动机，但并非其效果的唯一或最终解释。更稳妥的表述是：BN 重参数化并平滑/改变了优化几何，同时引入 batch 统计噪声和一定正则化效应。', 'FFF7E8')
source('L16 pp. 43–47；Ioffe & Szegedy (2015)。')

doc.add_heading('13. 把四条主线合在一起', level=1)
add_table(
    ['问题', '核心结论', '不能过度推断'],
    [
        ('表达', 'ReLU 网络能逼近广泛函数；深度可利用组合结构节省宽度。', '万能逼近不保证网络小或可训练。'),
        ('记忆', '较小的深网也能插值任意有限数据，甚至随机标签。', '零训练误差不等于零测试误差。'),
        ('优化', '反向传播高效算梯度；SGD 在非凸地形中提供可行搜索。', '凸 SGD 收敛率不能无条件套到深网。'),
        ('泛化', '范数、数据结构、算法隐式偏置比裸参数计数更关键。', '容量大并不自动意味着泛化好或坏。'),
        ('稳定', '初始化、残差、归一化与裁剪控制信号和梯度尺度。', '任何单一技巧都不是普适解释。'),
    ],
    [1.0, 3.3, 2.2],
)

doc.add_heading('13.1 学习与复习检查表', level=2)
bullet('能从核 SVM 的预测式解释“固定特征”与“学习特征”的区别。')
bullet('能写出任意深度前馈网络的 zˡ、aˡ 递推并检查矩阵维度。')
bullet('能从斜率变化推导一维分段仿射函数的 ReLU 表示。')
bullet('能说明万能逼近、有限样本记忆和 VC 维分别使用什么量词。')
bullet('能完整推导 δˡ、∂ℓ/∂Wˡ、∂ℓ/∂bˡ。')
bullet('能用方差守恒推导 ReLU 的 He 初始化 Var(w)=2/n。')
bullet('能从矩阵乘积解释梯度消失/爆炸。')
bullet('能写出 inverted dropout 与 BatchNorm 的前向、反向公式。')
bullet('能区分训练误差、经验风险、总体风险与三类误差。')

doc.add_heading('附录 A. 反向传播伪代码', level=1)
number('输入单个样本 (x,y)，令 a⁰=x。', restart=True)
number('前向：对 l=1,…,L，计算 zˡ=Wˡaˡ⁻¹+bˡ，aˡ=σˡ(zˡ)。')
number('输出误差：计算 δᴸ=∇_{zᴸ}ℓ(aᴸ,y)。')
number('反向：对 l=L−1,…,1，计算 δˡ=((Wˡ⁺¹)ᵀδˡ⁺¹)⊙σ′(zˡ)。')
number('梯度：对每层计算 G_Wˡ=δˡ(aˡ⁻¹)ᵀ，G_bˡ=δˡ。')
number('用 SGD/Adam 等更新参数；mini-batch 时对样本梯度求平均。')

doc.add_heading('附录 B. 来源对应与延伸阅读', level=1)
add_table(
    ['来源', '本讲义采用的相关部分'],
    [
        ('Sra, L15', '四种视角；网络结构；ReLU；非线性分离；万能逼近；深度分离；VC 维备注。'),
        ('Sra, L16', '有限样本记忆；SGD 实践；初始化；反向传播；梯度不稳定；Dropout；BatchNorm。'),
        ('Bach, Learning Theory from First Principles', 'Ch. 5 SGD；Ch. 9 单隐层网络、正齐次性、Rademacher 界、变差范数、逼近与核方法关系。'),
        ('Shalev-Shwartz & Ben-David', 'Ch. 20 前馈网络、表达能力、VC 维、训练困难与反向传播。'),
    ],
    [2.2, 4.3],
)

doc.add_heading('参考文献', level=2)
refs = [
    'Bach, F. Learning Theory from First Principles. MIT Press draft, 2025 version used here.',
    'Shalev-Shwartz, S. & Ben-David, S. Understanding Machine Learning: From Theory to Algorithms. Cambridge University Press, 2014.',
    'Cybenko, G. Approximation by superpositions of a sigmoidal function. Mathematics of Control, Signals and Systems, 1989.',
    'Leshno, M. et al. Multilayer feedforward networks with a nonpolynomial activation function can approximate any function. Neural Networks, 1993.',
    'He, K. et al. Delving Deep into Rectifiers. ICCV, 2015.',
    'Srivastava, N. et al. Dropout: A Simple Way to Prevent Neural Networks from Overfitting. JMLR, 2014.',
    'Ioffe, S. & Szegedy, C. Batch Normalization. ICML, 2015.',
    'Yun, C., Sra, S. & Jadbabaie, A. Small ReLU Networks Are Powerful Memorizers. NeurIPS, 2019.',
    'Bartlett, P. et al. Nearly-tight VC-dimension and pseudodimension bounds for piecewise linear neural networks. JMLR, 2019.',
]
for item in refs:
    bullet(item)

# Document metadata and final paragraph controls
doc.core_properties.title = '神经网络：表达、记忆与优化——L15–L16 整合讲义'
doc.core_properties.subject = 'Mathematical Foundations of Machine Learning'
doc.core_properties.author = 'OpenAI Codex (compiled from supplied course materials)'
doc.core_properties.keywords = 'neural networks, ReLU, backpropagation, SGD, memorization, batch normalization'

for p in doc.paragraphs:
    if p.style.name.startswith('Heading'):
        p.paragraph_format.keep_with_next = True
    if p.style.name == 'Equation':
        p.paragraph_format.keep_together = True

doc.save(DOCX)
print(DOCX.resolve())
